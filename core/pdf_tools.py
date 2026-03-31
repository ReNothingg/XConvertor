import os

from PyPDF2 import PdfReader, PdfWriter
from PIL import Image
import docx
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

try:
    import pypdfium2 as pdfium
except ImportError:
    pdfium = None


def _safe_close(resource):
    close = getattr(resource, 'close', None)
    if callable(close):
        close()

class PDFTools:
    def to_docx(self, pdf_path, docx_path):
        text = ''
        with open(pdf_path, 'rb') as f:
            reader = PdfReader(f)
            for page in reader.pages:
                text += (page.extract_text() or '') + '\n'
        
        doc = docx.Document()
        doc.add_paragraph(text)
        doc.save(docx_path)
        return docx_path

    def to_images(self, pdf_path, output_folder, fmt='jpeg'):
        if pdfium is None:
            raise RuntimeError(
                "Для конвертации PDF в изображения установите зависимость pypdfium2 "
                "через `pip install -r requirements.txt`."
            )

        os.makedirs(output_folder, exist_ok=True)
        output_ext = 'jpg' if fmt in {'jpg', 'jpeg'} else fmt
        image_format = 'JPEG' if output_ext == 'jpg' else output_ext.upper()
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]
        scale = 200 / 72
        paths = []
        pdf = pdfium.PdfDocument(pdf_path)

        try:
            for page_index in range(len(pdf)):
                page = pdf[page_index]
                bitmap = None
                image = None

                try:
                    bitmap = page.render(scale=scale)
                    image = bitmap.to_pil()

                    if output_ext == 'jpg' and image.mode != 'RGB':
                        converted = image.convert('RGB')
                        image.close()
                        image = converted

                    output_path = os.path.join(
                        output_folder,
                        f"{base_name}_page_{page_index + 1:03d}.{output_ext}",
                    )
                    image.save(output_path, format=image_format)
                    paths.append(output_path)
                finally:
                    if image is not None:
                        image.close()
                    _safe_close(bitmap)
                    _safe_close(page)
        finally:
            _safe_close(pdf)

        return paths

    def merge_pdfs(self, pdf_paths, output_path):
        writer = PdfWriter()
        for pdf_path in pdf_paths:
            reader = PdfReader(pdf_path)
            for page in reader.pages:
                writer.add_page(page)
        with open(output_path, 'wb') as f:
            writer.write(f)
        return output_path

    def split_pdf(self, pdf_path, output_folder):
        os.makedirs(output_folder, exist_ok=True)

        reader = PdfReader(pdf_path)
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]
        paths = []

        for page_number, page in enumerate(reader.pages, start=1):
            writer = PdfWriter()
            writer.add_page(page)

            output_path = os.path.join(output_folder, f"{base_name}_page_{page_number}.pdf")
            with open(output_path, 'wb') as f:
                writer.write(f)
            paths.append(output_path)

        return paths

    def images_to_pdf(self, image_paths, output_path):
        c = canvas.Canvas(output_path, pagesize=letter)
        width, height = letter
        for img_path in image_paths:
            img = Image.open(img_path)
            img_width, img_height = img.size
            aspect = img_height / float(img_width)
            
            display_width = width
            display_height = width * aspect
            
            if display_height > height:
                display_height = height
                display_width = height / aspect

            c.drawImage(img_path, 0, (height - display_height) / 2, width=display_width, height=display_height)
            c.showPage()
        c.save()
        return output_path
